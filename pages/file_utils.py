import io
import os
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx2pdf import convert
import platform
import pythoncom  # Added for COM initialization




def create_docx(cv_content):
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    main_color = RGBColor(30, 61, 89)
    styles = doc.styles
    styles.add_style('CVName', 1).font.size = Pt(18)
    styles['CVName'].font.bold = True
    styles.add_style('JobTitle', 1).font.size = Pt(14)
    styles['JobTitle'].font.italic = True
    styles.add_style('SectionTitle', 1).font.size = Pt(12)
    styles['SectionTitle'].font.bold = True
    styles['SectionTitle'].font.color.rgb = main_color
    styles.add_style('NormalText', 1).font.size = Pt(10)
    styles.add_style('Position', 1).font.size = Pt(11)
    styles['Position'].font.bold = True
    styles.add_style('Details', 1).font.size = Pt(10)
    styles['Details'].font.italic = True
    styles.add_style('Bullet', 1).font.size = Pt(10)
    styles['Bullet'].paragraph_format.left_indent = Cm(0.5)
    styles['Bullet'].paragraph_format.first_line_indent = Cm(-0.5)

    lines = cv_content.split('\n')
    current_section = None
    in_experience = False

    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        if i == 0:
            doc.add_paragraph(line, style='CVName')
        elif i == 1 and not line.upper() == line:
            doc.add_paragraph(line, style='JobTitle')
        elif line.upper() == line and len(line) < 40:
            current_section = line
            in_experience = "EXPÉRIENCE" in current_section or "EXPERIENCE" in current_section
            p = doc.add_paragraph(line, style='SectionTitle')
            p.border_bottom = True
            p.border_bottom_color = main_color
            p.border_bottom_width = Pt(1)
        elif current_section:
            if in_experience and (line.startswith('•') or line.startswith('-')):
                doc.add_paragraph(line.lstrip('•').lstrip('-').strip(), style='Bullet')
            elif in_experience and '|' in line:
                p = doc.add_paragraph('', style='Position')
                parts = line.split('|')
                for j, part in enumerate(parts):
                    part = part.strip()
                    run = p.add_run(part)
                    if j < 2:  # Bold job title and company
                        run.bold = True
                    if j < len(parts) - 1:
                        run.add_text(' | ')
            elif in_experience:
                doc.add_paragraph(line, style='Details')
            elif line.startswith('•') or line.startswith('-'):
                doc.add_paragraph(line.lstrip('•').lstrip('-').strip(), style='Bullet')
            else:
                doc.add_paragraph(line, style='NormalText')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Convertir DOCX en PDF
def create_pdf_from_docx(docx_buffer):
    try:
        # Initialize COM on Windows with multi-threaded support
        if platform.system() == "Windows":
            try:
                pythoncom.CoInitializeEx(pythoncom.COINIT_MULTITHREADED)
            except pythoncom.com_error as e:
               print(f"⚠️ COM initialization failed: {e}. Attempting to proceed.")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
            temp_docx.write(docx_buffer.getvalue())
            temp_docx_path = temp_docx.name

        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf_path = temp_pdf.name

        # Perform DOCX to PDF conversion
        convert(temp_docx_path, temp_pdf_path)

        with open(temp_pdf_path, "rb") as f:
            pdf_buffer = io.BytesIO(f.read())

        os.unlink(temp_docx_path)
        os.unlink(temp_pdf_path)

        pdf_buffer.seek(0)
        return pdf_buffer
    except Exception as e:
        return None
    finally:
        # Uninitialize COM on Windows
        if platform.system() == "Windows":
            try:
                pythoncom.CoUninitialize()
            except:
                pass
